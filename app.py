import streamlit as st
import pandas as pd
import datetime
import os
import re
import json
from streamlit_autorefresh import st_autorefresh 
from rapidfuzz import process, utils
from streamlit_js_eval import streamlit_js_eval
import pytz
from garminconnect import Garmin
import requests
from streamlit_oauth import OAuth2Component
import webbrowser
import io
from docx import Document
from docx.shared import Pt

# ============================================
# --- 1. TIMEZONE & REFRESH ---
# ============================================

user_tz_name = streamlit_js_eval(js_expressions='Intl.DateTimeFormat().resolvedOptions().timeZone', key='tz')
user_tz = pytz.timezone(user_tz_name) if user_tz_name else pytz.timezone("UTC")

# Refresh the app every 30 seconds to update the "Live" tally
st_autorefresh(interval=30000, key="daterefresh")

# ============================================
# --- 2. GLOBAL STATE INITIALIZATION ---
# ============================================

if "bmr" not in st.session_state:
    # Physical Specs (Defaults)
    st.session_state["gender"] = "Male"
    st.session_state["age"] = 25
    st.session_state["height_in"] = 70.0 
    st.session_state["weight_lbs"] = 190.0
    st.session_state["history"] = []
    
    # Gear & Sync States (Initializes to 0 so HUD doesn't crash)
    st.session_state["sync_active_burn"] = 0
    st.session_state["csv_active_burn"] = 0
    st.session_state["oura_results"] = {"active_cals": 0, "readiness": 80, "sleep": 80}
    
    # Calculate BMR (Mifflin-St Jeor)
    w_kg = st.session_state["weight_lbs"] * 0.453592
    h_cm = st.session_state["height_in"] * 2.54
    a = st.session_state["age"]
    
    if st.session_state["gender"] == "Male":
        st.session_state["bmr"] = (10 * w_kg) + (6.25 * h_cm) - (5 * a) + 5
    else:
        st.session_state["bmr"] = (10 * w_kg) + (6.25 * h_cm) - (5 * a) - 161

# Local variables for easy use in logic below

# This fixes the "NameError" in your UI by pulling the values out of storage
bmr = st.session_state.get("bmr", 2000.0)
eaten = st.session_state.get("eaten", 0.0)
req_deficit = st.session_state.get("req_deficit", 500.0)

# Initialize device states so they don't "ghost" in from previous runs
# Numerical trackers start at 0
for key in ["sync_active_burn", "fitbit_active_burn", "suunto_active_burn", "whoop_active_burn", "manual_active_burn"]:
    if key not in st.session_state:
        st.session_state[key] = 0

# Oura results must be a dictionary to prevent crashes
if "oura_results" not in st.session_state or st.session_state["oura_results"] == 0:
    st.session_state["oura_results"] = {}

# ============================================
# --- 3. TACTICAL GEAR INTEGRATION LAYER ---
# ============================================

from garminconnect import Garmin
import datetime

def fetch_garmin_active_burn(email, password):
    """Logs into Garmin and fetches today's active calories."""
    try:
        # Initialize client
        client = Garmin(email, password)
        
        # Attempt Login
        client.login()
        
        # Get Today's Summary
        today = datetime.date.today().isoformat()
        stats = client.get_user_summary(today)
        
        # Extract active calories - Checking both possible keys
        active_burn = stats.get("activeCalories") or stats.get("activeKilocalories") or 0
        
        return int(active_burn)

    except Exception as e:
        # This will catch wrong passwords or connection issues
        raise Exception(f"Authentication Failed: {str(e)}")

def render_oura_gear_ui():
    """Fetch credentials only when the UI is rendered."""
    # Pull secrets locally inside the function
    try:
        c_id = st.secrets["oura"]["client_id"]
        c_secret = st.secrets["oura"]["client_secret"]
        r_uri = st.secrets["oura"]["redirect_uri"]
    except KeyError:
        st.error("Oura secrets not found in .streamlit/secrets.toml")
        return

    auth_url = "https://cloud.ouraring.com/oauth/authorize"
    token_url = "https://api.ouraring.com/oauth/token"

    oauth2 = OAuth2Component(c_id, c_secret, auth_url, token_url, token_url, token_url)

    if "oura_token" not in st.session_state:
        st.info("Oura Ring not connected.")
        result = oauth2.authorize_button(
            name="💍 Connect Oura Ring",
            scope="daily personal", 
            redirect_uri=r_uri,
        )
        if result:
            st.session_state["oura_token"] = result["token"]
            st.rerun()
    else:
        st.success("Oura Connected ✅")
        if st.button("Disconnect Oura"):
            del st.session_state["oura_token"]
            st.rerun()

def fetch_oura_v2_data(token):
    import requests
    from datetime import datetime, timedelta
    
    headers = {'Authorization': f'Bearer {token}'}
    # Look back 3 days to catch the sync regardless of timezone
    start = (datetime.now() - timedelta(days=3)).strftime('%Y-%m-%d')
    url = f'https://api.ouraring.com/v2/usercollection/daily_activity?start_date={start}'
    
    try:
        res = requests.get(url, headers=headers)
        if res.status_code == 200:
            data = res.json().get('data', [])
            if not data:
                return {"active_cals": 0}
            
            # Sort by date and grab the very latest calories found
            valid_entries = [d for d in data if d.get('active_calories', 0) > 0]
            if valid_entries:
                # Returns the most recent non-zero calorie count
                return {"active_cals": valid_entries[-1]['active_calories']}
        return {"active_cals": 0}
    except:
        return {"active_cals": 0}
    
    # Use standard ISO format
    today_str = datetime.date.today().strftime('%Y-%m-%d')
    
    # Note: We query specifically for today to avoid the historical '80' baseline
    url = f"https://api.ouraring.com/v2/usercollection/daily_readiness?start_date={today_str}&end_date={today_str}"
    
    try:
        response = requests.get(url, headers=headers)
        # 💡 If this prints 401, your token expired. If 403, your scope is wrong.
        if response.status_code != 200:
            st.error(f"Oura API Error {response.status_code}: {response.text}")
            return {}

        data = response.json().get('data', [])
        
        if data:
            latest = data[-1]
            return {
                "readiness_score": latest.get('score', 0),
                "date_synced": latest.get('day'),
                "contributors": latest.get('contributors', {})
            }
        else:
            # This happens if Oura hasn't computed Today's score yet
            return {"readiness_score": "Pending", "date_synced": today_str}
            
    except Exception as e:
        st.error(f"Sync Failed: {e}")
        return {}

def parse_universal_health_csv(uploaded_file):
    try:
        df = pd.read_csv(uploaded_file)
        data = {"active_burn": 0, "readiness": 80, "sleep": 80}
        for col in df.columns:
            c = col.lower()
            if ('active' in c or 'kcal' in c) and ('calorie' in c or 'energy' in c):
                data["active_burn"] = int(df[col].sum())
            if 'readiness' in c or 'recovery' in c:
                data["readiness"] = int(df[col].iloc[-1])
            if 'sleep' in c and 'score' in c:
                data["sleep"] = int(df[col].iloc[-1])
        return data
    except: return None

# ==========================================
# --- THE LIVE HUD MATH (The Engine) ---
# ==========================================

# 1. Calculate BMR Tally based on time of day
now = datetime.datetime.now()
seconds_since_midnight = (now.hour * 3600) + (now.minute * 60) + now.second
percent_of_day = seconds_since_midnight / 86400
live_bmr = int(bmr * percent_of_day)

# 2. Consolidate Active Burn (Priority: Garmin > CSV > Oura > Manual)
garmin_val = st.session_state.get("sync_active_burn", 0)
csv_val = st.session_state.get("csv_active_burn", 0)
oura_val = st.session_state.get("oura_results", {}).get("active_cals", 0)
sensor_active = max(garmin_val, csv_val, oura_val)

# 3. Fallback to manual if no sensors are synced
manual_active = 0
if st.session_state.get("history"):
    today_date = datetime.date.today()
    manual_active = sum(e["Value"] for e in st.session_state.history 
                        if e["Date"] == today_date and e["Category"] in ["Active", "Exercise"])

# 4. FINAL TALLY
final_active_today = sensor_active if sensor_active > 0 else manual_active
total_live_burn = live_bmr + final_active_today

# ==========================================
# --- THE BRIDGE (THE NAMES THE UI NEEDS) ---
# ==========================================

# Place these lines here to stop the "NameError" crashes
active_today = final_active_today
bmr = st.session_state.get("bmr", 2000.0)
eaten = st.session_state.get("eaten", 0.0)
req_deficit = st.session_state.get("req_deficit", 500.0)

# ============================================
# --- PERSISTENCE LOGIC (Add this back in) ---
# ============================================

STRATEGY_FILE = "mission_strategy.json"

def save_strategy(strategy_text, target_desc, start_date):
    data = {
        "plan": strategy_text,
        "target": target_desc,
        "generated_on": str(start_date)
    }
    with open(STRATEGY_FILE, "w") as f:
        json.dump(data, f)

def load_strategy():
    if os.path.exists(STRATEGY_FILE):
        try:
            with open(STRATEGY_FILE, "r") as f:
                return json.load(f)
        except:
            return None
    return None

DB_FILE = "gym_locker.json" 

def save_gyms(data):
    with open(DB_FILE, "w") as f:
        json.dump(data, f)

def load_gyms():
    if os.path.exists(DB_FILE):
        try:
            with open(DB_FILE, "r") as f:
                return json.load(f)
        except:
            return None
    return None

# Re-link the saved data to your session state
saved_data = load_gyms()
if saved_data:
    st.session_state["gym_profiles"] = saved_data

STORE_FILE = "target_stores.json"

def save_stores(data):
    import json
    import os
    with open(STORE_FILE, "w") as f:
        json.dump(data, f)

def load_stores():
    import json
    import os
    if os.path.exists(STORE_FILE):
        try:
            with open(STORE_FILE, "r") as f:
                return json.load(f)
        except:
            return None
    return None

# Initialize/Load Stores into Session State immediately after defining the functions
if "custom_stores" not in st.session_state:
    saved_stores = load_stores()
    if saved_stores:
        st.session_state["custom_stores"] = saved_stores
    else:
        st.session_state["custom_stores"] = ["NEX / Commissary", "Costco", "Sam's Club", "Walmart", "Aldi"]

# ============================================
# --- INITIALIZE STRATEGY IN SESSION STATE ---
# ============================================

# Add this right after your gym_profiles initialization
if "long_term_plan" not in st.session_state:
    saved_strat = load_strategy()
    if saved_strat:
        st.session_state["long_term_plan"] = saved_strat["plan"]
        st.session_state["perf_target"] = saved_strat["target"]
        st.session_state["strategy_generated_on"] = datetime.datetime.strptime(saved_strat["generated_on"], '%Y-%m-%d').date()

if "account_created" not in st.session_state:
    # Defaults to 30 days ago if new, or you can set a specific date
    st.session_state.account_created = datetime.date.today() - datetime.timedelta(days=30)

# --- INITIAL SETUP & TACTICAL THEME ---
st.set_page_config(
    page_title="Navy PFA Pro", 
    layout="wide", 
    page_icon="⚓",
    initial_sidebar_state="expanded" # Keeps it open on load, but allows user to close
)

st.markdown("""
    <style>
        /* 1. TOP SPACING & RESPONSIVE CONTAINER */
        .block-container {
            padding-top: 1rem !important;
            padding-bottom: 0rem;
            /* Remove manual left/right padding to allow the 'wide' layout to breathe */
            margin-top: -20px; 
        }
        
        /* Hide decoration but keep the functional header for the sidebar toggle */
        [data-testid="stHeader"] {
            background: rgba(0,0,0,0) !important;
        }
        footer {visibility: hidden;}

        /* 2. SIDEBAR TACTICAL LOOK (Without breaking responsiveness) */
        [data-testid="stSidebar"] {
            background-color: #0E1117 !important;
            border-right: 1px solid rgba(0, 229, 255, 0.2);
        }

        /* Make the toggle button (the hamburger/arrow) Neon Cyan */
        [data-testid="stSidebarCollapseButton"] {
            color: #00E5FF !important;
        }

        /* 3. FORCE DARK THEME & TACTICAL COLORS */
        [data-testid="stAppViewContainer"], .main {
            background-color: #0E1117 !important;
            color: #E0E0E0 !important;
        }

        h1, h2, h3 {
            color: #00E5FF !important;
            text-transform: uppercase;
            letter-spacing: 2px;
            border-left: 5px solid #00E5FF;
            padding-left: 15px;
            margin-top: 5px !important;
        }

        /* 4. INPUTS, BUTTONS & METRICS */
        input, select, textarea, [data-baseweb="select"] {
            background-color: #1A1C23 !important;
            color: white !important;
            border: 1px solid rgba(0, 229, 255, 0.3) !important;
        }

        .stButton>button {
            background-color: transparent !important;
            color: #00E5FF !important;
            border: 1px solid #00E5FF !important;
            transition: all 0.3s ease;
        }

        .stButton>button:hover {
            background-color: #00E5FF !important;
            color: #000000 !important;
            box-shadow: 0 0 10px #00E5FF;
        }

        [data-testid="stMetricValue"] {
            color: #00E5FF !important;
            font-family: 'Courier New', monospace;
        }

        /* 5. MULTISELECT CHIP FIX */
        /* Make the internal search box transparent so it doesn't overlap text */
        [data-baseweb="select"] input {
            background-color: transparent !important;
            border: none !important;
        }
        
        /* Change the selected chips from default red to Tactical Cyan */
        [data-baseweb="tag"] {
            background-color: rgba(0, 229, 255, 0.1) !important;
            border: 1px solid #00E5FF !important;
            color: #00E5FF !important;
        }

        /* CUSTOM SCROLLBAR */
        ::-webkit-scrollbar { width: 8px; }
        ::-webkit-scrollbar-thumb {
            background: #1A1C23;
            border: 1px solid #00E5FF;
        }
    </style>
""", unsafe_allow_html=True)

# --- INITIALIZE GEMINI CLIENT ---
from google import genai

# 1. Check for the section first to avoid a crash
AI_READY = False
AI_ERROR = ""

if "gemini" in st.secrets:
    try:
        # 1. Standardize credentials
        g_api_key = st.secrets["gemini"]["api_key"]
        
        # 2. Initialize the Client
        client = genai.Client(api_key=g_api_key)
        
        # 3. Set your 2026 model (Gemini 2.5 Flash is the current workhorse)
        MODEL_ID = st.secrets["gemini"].get("model", "gemini-2.5-flash")
        
        AI_READY = True
    except Exception as e:
        AI_READY = False
        AI_ERROR = str(e)
        MODEL_ID = "gemini-2.5-flash" # Fallback string to prevent crashes
else:
    AI_READY = False
    AI_ERROR = "Configuration Missing"
    MODEL_ID = "gemini-2.5-flash"

# Initialize/Load Gym Profiles
if "gym_profiles" not in st.session_state:
    saved_data = load_gyms()
    if saved_data:
        st.session_state["gym_profiles"] = saved_data
    else:
        # 🚨 FIX: These strings now perfectly match the standard_gear list below
        st.session_state["gym_profiles"] = {
            "🏠 Home Gym": ["Bodyweight", "Dumbbells", "Pull-up Bar"],
            "⚓ Ship Gym (Afloat)": ["Bodyweight", "Kettlebells", "Pull-up Bar"],
            "🏋️ Base Gym (Main)": ["Barbell & Plates", "Rowing Machine", "Stairmaster"]
        }

# Initialize/Load Stores
if "custom_stores" not in st.session_state:
    saved_stores = load_stores()
    if saved_stores:
        st.session_state["custom_stores"] = saved_stores
    else:
        st.session_state["custom_stores"] = ["NEX / Commissary", "Costco", "Sam's Club", "Walmart", "Aldi"]

# ==========================================
# --- DATA FUNCTIONS ---
# ==========================================

def get_navy_tier(score):
    """Maps 0-100 score to official Navy Tiers."""
    if score >= 90:
        if score >= 95: return "Outstanding (High)", "#00E5FF"
        if score >= 92: return "Outstanding (Medium)", "#00B8D4"
        return "Outstanding (Low)", "#0097A7"
    elif score >= 75:
        if score >= 85: return "Excellent (High)", "#00C853"
        if score >= 80: return "Excellent (Medium)", "#2E7D32"
        return "Excellent (Low)", "#388E3C"
    elif score >= 60:
        if score >= 70: return "Good (High)", "#FFAB00"
        if score >= 65: return "Good (Medium)", "#FF8F00"
        return "Good (Low)", "#EF6C00"
    elif score >= 45:
        return "Satisfactory", "#D32F2F"
    else:
        return "Probationary/Fail", "#757575"

def calculate_score(val, event_type, age, gender):
    """Calculates point values for PRT events."""
    if event_type == "Pushups":
        ref = 60 if gender == "Male" else 35
        return min(100, int((val / ref) * 80))
    elif event_type == "Plank":
        return min(100, int((val / 180) * 90))
    elif event_type == "1.5 Mile Run":
        ref_time = 630 if gender == "Male" else 750
        return min(100, int((ref_time / val) * 95))
    elif event_type == "Stationary Bike":
        ref_cals = 200 if gender == "Male" else 160
        return min(100, int((val / ref_cals) * 90))
    else:
        return 75

def update_global_deadline():
    # Safety check: Only update if the key exists in session state
    if "master_date_input" in st.session_state:
        st.session_state.mission_deadline = st.session_state.master_date_input
def load_data():
    if os.path.exists('fitness_log.csv'):
        try:
            df = pd.read_csv('fitness_log.csv')
            df['date'] = df['date'].astype(str)
            return df
        except:
            return pd.DataFrame(columns=['date', 'type', 'description', 'calories'])
    return pd.DataFrame(columns=['date', 'type', 'description', 'calories'])

def save_entry(category, activity, value, date_obj=None):
    """The ONE function to rule them all. Handles HUD, CSV, and Burn logic."""
    # 1. Default to today if no date is provided
    if date_obj is None:
        date_obj = datetime.date.today()
    
    # 2. Type Enforcement & Safety Clamping
    try:
        clean_value = float(value)
    except (ValueError, TypeError):
        return st.error(f"Invalid numeric value: {value}")
    
    safe_value = max(0.0, min(clean_value, 5000.0))
    
    # 3. Handle 'Burn' math (CSV needs negative numbers for Exercise/Passive)
    # This keeps your old spreadsheet math consistent!
    csv_value = -abs(safe_value) if category in ["Exercise", "Passive", "Active"] else abs(safe_value)

    # 4. Save to Session State (for the Live HUD)
    if "history" not in st.session_state:
        st.session_state.history = []
    
    st.session_state.history.append({
        "Date": date_obj,
        "Category": category,
        "Activity": str(activity),
        "Value": safe_value # HUD likes positive numbers
    })

    # 5. Save to CSV (The 'Permanent' log)
    file_exists = os.path.exists('fitness_log.csv')
    new_row = pd.DataFrame([[str(date_obj), category, activity, csv_value]], 
                            columns=['date', 'type', 'description', 'calories'])
    new_row.to_csv('fitness_log.csv', mode='a', header=not file_exists, index=False)

def load_pantry():
    if os.path.exists('pantry.csv'):
        return pd.read_csv('pantry.csv')
    return pd.DataFrame(columns=['item', 'quantity'])

def update_pantry(item, qty, action="add"):
    pantry = load_pantry()
    item = item.lower().strip()
    if item in pantry['item'].values:
        if action == "add":
            pantry.loc[pantry['item'] == item, 'quantity'] += qty
        else:
            pantry.loc[pantry['item'] == item, 'quantity'] -= qty
    elif action == "add":
        new_item = pd.DataFrame([[item, qty]], columns=['item', 'quantity'])
        pantry = pd.concat([pantry, new_item], ignore_index=True)
    pantry = pantry[pantry['quantity'] > 0]
    pantry.to_csv('pantry.csv', index=False)

def smart_sync_garmin(email, password):
    """Dynamically determines sync window based on last recorded log."""
    try:
        client = Garmin(email, password)
        client.login()
        
        today = datetime.date.today()
        
        # 1. Determine the Last Sync Date from existing history
        history_df = pd.DataFrame(st.session_state.get("history", []))
        last_sync_date = None
        
        if not history_df.empty:
            # Look for previous Garmin entries specifically
            garmin_logs = history_df[history_df['Activity'].str.contains("Garmin", na=False)]
            if not garmin_logs.empty:
                last_sync_date = pd.to_datetime(garmin_logs['Date']).max().date()

        # 2. Calculate the "Start Date" based on your 14-day rule
        if last_sync_date is None:
            # Initial Connection: Go back 14 days
            start_date = today - datetime.timedelta(days=14)
            sync_type = "Initial Discovery"
        else:
            days_since_sync = (today - last_sync_date).days
            if days_since_sync > 14:
                # Gap > 14 days: Reset to 14-day window
                start_date = today - datetime.timedelta(days=14)
                sync_type = "Re-established Link (14-day reset)"
            else:
                # Regular Sync: Start from the day after the last log
                start_date = last_sync_date + datetime.timedelta(days=1)
                sync_type = "Differential Update"

        # 3. Perform the Sync
        delta = (today - start_date).days
        synced_count = 0
        
        # If we are already up to date, skip the loop
        if delta < 0:
            return 0, "System already synchronized."

        for i in range(delta + 1):
            target_date = start_date + datetime.timedelta(days=i)
            date_str = target_date.isoformat()
            
            stats = client.get_user_summary(date_str)
            burn = stats.get("activeCalories") or stats.get("activeKilocalories") or 0
            
            if burn > 0:
                # Add to local history (Match your existing save_entry structure)
                new_entry = {
                    "Date": target_date.strftime("%Y-%m-%d"),
                    "Category": "Exercise",
                    "Activity": f"Garmin Sync",
                    "Value": burn,
                    "Unit": "kcal"
                }
                st.session_state.history.append(new_entry)
                synced_count += 1
        
        # Save to CSV so it persists for the Analysis Tab
        pd.DataFrame(st.session_state.history).to_csv("fitness_log.csv", index=False)
        
        return synced_count, sync_type

    except Exception as e:
        raise Exception(f"Sync Protocol Aborted: {str(e)}")

def sync_oura_callback(token):
    # This runs BEFORE the page reruns
    res = fetch_oura_v2_data(token)
    st.session_state["oura_active_burn"] = res.get("active_cals", 0)
    # Give the user a toast notification
    if st.session_state["oura_active_burn"] > 0:
        st.toast(f"💍 Oura Synced: {st.session_state['oura_active_burn']} kcal!")

def global_safeguard():
    """Hard-limits all session data to humanly possible ranges."""
    # 1. Physical Limits
    if "age" in st.session_state:
        st.session_state.age = int(max(17, min(st.session_state.age, 65)))
    if "height" in st.session_state:
        st.session_state.height = float(max(48, min(st.session_state.height, 96)))
    if "current_weight" in st.session_state:
        st.session_state.current_weight = int(max(90, min(st.session_state.current_weight, 500)))

    # 2. History Integrity
    if "history" in st.session_state:
        # Prevent memory bloat (keep last 200 entries)
        if len(st.session_state.history) > 200:
            st.session_state.history = st.session_state.history[-200:]
        
        # Remove corrupted entries (None values or extreme numbers)
        st.session_state.history = [
            e for e in st.session_state.history 
            if isinstance(e.get("Value"), (int, float)) and 0 <= e["Value"] <= 5000
        ]

def create_word_doc(content):
    """Generates a formal Word Doc with a Navy-style header"""
    doc = Document()
    
    # Standard Naval Letterhead Style Header
    header = doc.sections[0].header
    htab = header.paragraphs[0]
    htab.text = "NAVY PFA PRO\tTACTICAL MISSION BRIEFING\tCONFIDENTIAL"
    htab.style = 'Header'
    
    doc.add_heading('MISSION PERFORMANCE PLAN', 0)
    doc.add_paragraph(f"DATE: {datetime.date.today().strftime('%d %b %Y').upper()}")
    doc.add_paragraph("-" * 50)

    for line in content.split('\n'):
        line = line.strip()
        if line.startswith('###'):
            doc.add_heading(line.replace('###', ''), level=2)
        elif line.startswith('##'):
            doc.add_heading(line.replace('##', ''), level=1)
        elif line:
            p = doc.add_paragraph(line)
            p.style.font.name = 'Courier New' # Tactical mono font
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_excel_plan(content):
    """Converts AI output into a row-based Excel tracker"""
    lines = [line.strip() for line in content.split('\n') if line.strip()]
    df = pd.DataFrame(lines, columns=["Mission Directives"])
    
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='7-Day-Tactical')
    return bio.getvalue()

# Execute immediately
global_safeguard()

@st.dialog("Confirm Deletion")
def confirm_delete_dialog(indices):
    st.warning(f"⚠️ This will permanently delete {len(indices)} entries. Proceed?")
    if st.button("🔥 Confirm Bulk Purge", type="primary", use_container_width=True):
        df = load_data()
        df = df.drop(indices)
        df.to_csv('fitness_log.csv', index=False)
        st.session_state['df_key'] += 1 
        st.rerun()

# ==========================================================
# --- SIDEBAR: MISSION COMMAND, COMMS-LINK & MAINTENANCE ---
# ==========================================================

# 1. Define how we identify "Active Burn" for EVERY potential device
DEVICE_MAP = {
    "Garmin": st.session_state.get("sync_active_burn", 0),
    "Oura": st.session_state.get("oura_active_burn", 0),
    "Fitbit": st.session_state.get("fitbit_active_burn", 0),
    "Suunto": st.session_state.get("suunto_active_burn", 0),
    "Whoop": st.session_state.get("whoop_active_burn", 0),
    "Strava": st.session_state.get("strava_active_burn", 0),
    "Manual": st.session_state.get("manual_active_burn", 0)
}

# 2. STRICT FILTER: Only include devices that actually have data > 0
connected_devices = [
    name for name, val in DEVICE_MAP.items() 
    if isinstance(val, (int, float)) and val > 0
]

with st.sidebar:
    st.header("🚢 Mission Control")
    
    # --- STRATEGIC OVERVIEW ---
    deadline = st.session_state.get("mission_deadline", datetime.date.today())
    days_to_go = (deadline - datetime.date.today()).days
    
    st.metric(label="Mission Deadline", value=deadline.strftime("%d %b %Y"))
    st.caption(f"🏁 {days_to_go} Days Remaining")
    st.divider()    

    # --- DAILY BRIEFING ---
    st.subheader("📋 Today's Brief")
    if "daily_evo" in st.session_state:
        st.info("**Workout Evolution Active**")
    else:
        st.caption("No workout generated for today.")

    if "active_nut_plan" in st.session_state:
        st.success("**Nutrition Plan Active**")
    else:
        st.caption("No meal plan generated for today.")

    current_tier = st.session_state.get("current_tier", "Unknown")
    st.caption(f"Current Tier: **{current_tier}**")
    st.divider()
    
    # --- COMMS-LINK (GEAR & AI SYNC) ---
    st.subheader("📡 Comms-Link")

    if AI_READY:
        st.success("🟢 **Gemini AI:** Online", icon="🧠")
    else:
        st.error("🔴 **Gemini AI:** Offline", icon="⚠️")

    # Primary Tracker Selector
    if len(connected_devices) >= 2:
        st.markdown("---")
        primary_source = st.radio(
            "🎯 Primary Energy Source",
            options=connected_devices,
            help="Multiple sources detected. Choose which one leads the Active Burn metric.",
            horizontal=True,
            key="user_selected_primary"
        )
    else:
        primary_source = connected_devices[0] if connected_devices else "Manual"

    st.markdown("---")

    gear_source = st.selectbox(
        "Select External Gear",
        ["Standby", "Garmin Connect", "Oura Ring", "Fitbit", "Strava", "Whoop", "Suunto"],
        key="primary_gear_selector_v16" 
    )

    # --- GEAR UPLINK LOGIC ---
    if gear_source == "Garmin Connect":
        st.markdown("### ⚓ GARMIN SMART SYNC")
        
        # 1. Local Mode (Secrets Exist)
        if "garmin" in st.secrets:
            st.success("✅ Secure Vault credentials loaded.")
            u_email = st.secrets["garmin"]["email"]
            u_pass = st.secrets["garmin"]["password"]
            
        # 2. Cloud/Guest Mode (No Secrets)
        else:
            st.info("Guest Mode: Enter your Garmin credentials to sync.")
            u_email = st.text_input("Garmin Email", key="guest_g_user")
            u_pass = st.text_input("Garmin Password", type="password", key="guest_g_pass")
            
        if st.button("📡 Execute Tactical Sync", use_container_width=True, type="primary"):
            if not u_email or not u_pass:
                st.error("Credentials Required.")
            else:
                with st.spinner("Analyzing mission logs and syncing..."):
                    try:
                        count, s_type = smart_sync_garmin(u_email, u_pass)
                        if count > 0:
                            st.success(f"✅ {s_type}: {count} days ingested.")
                            # This updates the 'Active Burn' metric for the HUD immediately
                            st.session_state["sync_active_burn"] = count 
                        else:
                            st.info("💡 All systems synchronized. No new data found.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Sync Error: {e}")

    elif gear_source == "Oura Ring":
        st.markdown("### 💍 OURA LINK")
        
        # 1. Local Mode (Secrets Exist)
        if "oura" in st.secrets:
            st.success("✅ Secure Vault token loaded.")
            active_token = st.secrets["oura"]["token"]
            
        # 2. Cloud/Guest Mode (No Secrets)
        else:
            st.info("Guest Mode: Enter your Oura Personal Access Token.")
            active_token = st.text_input("Manual Token Override", type="password", key="oura_man_token_final")
            
        if st.button("🔄 Sync Oura", use_container_width=True, type="primary"):
            if not active_token:
                st.error("Token Required.")
            else:
                with st.spinner("Uplinking Oura telemetry..."):
                    try:
                        res = fetch_oura_v2_data(active_token)
                        # SAVE THE ENTIRE RESULT OBJECT
                        st.session_state["oura_results"] = res 
                        # ALSO SAVE THE BURN SPECIFICALLY FOR THE TRACKER
                        st.session_state["oura_active_burn"] = res.get("active_cals", 0)
                        st.success("Successfully Synced!")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Sync Error: {e}")

    elif gear_source in ["Fitbit", "Strava", "Whoop", "Suunto"]:
        st.info(f"Authorization for {gear_source} required.")
        if st.button(f"🔗 Authorize {gear_source}", use_container_width=True):
            st.session_state[f"{gear_source.lower()}_active_burn"] = 500
            st.rerun()

    # --- UNIVERSAL CONNECTION STATUS & DISCONNECT ---
    st.divider()
    active_device = None
    current_val = 0
    device_status_map = {"Garmin Connect": "sync_active_burn", "Oura Ring": "oura_active_burn"}
    
    for device, key in device_status_map.items():
        if st.session_state.get(key, 0) > 0:
            active_device = device
            current_val = st.session_state[key]
            break

    if active_device:
        st.write(f"**{active_device} Status**")
        st.success(f"🟢 Connected")
        st.info(f"⚡ Active Burn: {current_val} kcal")
        if st.button("❌ Disconnect Device", use_container_width=True):
            for key in device_status_map.values():
                st.session_state[key] = 0
            st.rerun()
    else:
        st.caption("⚪ System Standby: No devices linked")

    # --- SYSTEM MAINTENANCE (Always Visible) ---
    st.write("<br>", unsafe_allow_html=True) 
    with st.expander("🛠️ System Maintenance"):
        st.write("Manage Local Data Files")
        confirm_reset = st.checkbox("Confirm Strategy Reset", key="sidebar_strategy_reset_check")
        if st.button("Reset Mission Strategy", use_container_width=True, disabled=not confirm_reset):
            if os.path.exists("mission_strategy.json"):
                os.remove("mission_strategy.json")
            st.success("Strategy Deleted.")
            st.rerun()
        
        st.markdown("---")
        if st.button("Clear Gym Locker", use_container_width=True):
            if os.path.exists("gym_locker.json"):
                os.remove("gym_locker.json")
                st.rerun()
        
        if st.button("Purge All Logs", type="primary", use_container_width=True):
            if os.path.exists("fitness_log.csv"): # Updated to match your main saving logic
                os.remove("fitness_log.csv")
                st.session_state.history = []
                st.rerun()

    # --- SYSTEM DIAGNOSTICS (Always Visible) ---
    with st.expander("🔍 System Diagnostics", expanded=False):
        st.write("### 🔑 Secret Vault Check")
        try:
            # List Top Level Secrets
            all_sec_keys = list(st.secrets.keys())
            st.write(f"**Loaded Sections:** `{all_sec_keys}`")
            
            # Drill down into specific app sections if they exist
            for section in ["garmin", "oura", "gemini"]:
                if section in st.secrets:
                    sub_keys = list(st.secrets[section].keys())
                    st.write(f"**{section.capitalize()} Keys Found:** `{sub_keys}`")
        except Exception as e:
            st.error("Could not read secrets. check .streamlit/secrets.toml formatting.")

        st.divider()
        st.write("### 📂 File System")
        st.write(f"**App Path:** `{os.getcwd()}`")
        st.write(f"**Config Folder:** `{'✅ Found' if os.path.exists('.streamlit') else '❌ Missing'}`")
        
        if st.button("Clear App Cache"):
            st.cache_data.clear()
            st.cache_resource.clear()
            st.toast("Cache Cleared!")

    # --- COPYRIGHT FOOTER ---
    st.write("<br><br>", unsafe_allow_html=True) # Pushes the footer down slightly
    st.markdown(
        """
        <div style='text-align: center; color: rgba(0, 229, 255, 0.4); font-size: 0.8em; font-family: "Courier New", monospace;'>
            © 2026 PFA-Pro.<br>All Rights Reserved.
        </div>
        """, 
        unsafe_allow_html=True
    )

# ==========================================
# --- TACTICAL HUD ---
# ==========================================

# --- UNIVERSAL HUD BRIDGE ---
# Look up the burn value based on the selected source name
active_today = DEVICE_MAP.get(primary_source, 0)

st.subheader("⚓ Real-Time Mission Energy Status")

# Create 4 columns for the full breakdown
col1, col2, col3, col4 = st.columns(4)

# Set an icon based on the source
icons = {"Garmin": "⚓", "Oura": "💍", "Fitbit": "⌚", "Suunto": "⛰️", "Manual": "📝"}
active_icon = icons.get(primary_source, "⚡")

with col1:
    st.metric(
        label="🧘 Passive Burn", 
        value=f"{live_bmr} kcal", 
        help="Estimated calories burned by your body so far today (BMR)."
    )

with col2:
    st.metric(
        label="🏃 Active Burn", 
        value=f"{active_today} kcal", 
        help="Sum of all exercise/active entries logged today."
    )

with col3:
    st.metric(
        label="🍽️ Consumption", 
        value=f"{int(eaten)} kcal", 
        help="Total calories consumed today."
    )

with col4:
    # Calculation: What you've eaten vs. what you've burned (Passive + Active)
    # Negative = Deficit (Weight Loss), Positive = Surplus (Weight Gain)
    current_balance = eaten - (live_bmr + active_today)
    
    # Logic to change the label based on the state
    status_label = "📈 Surplus" if current_balance > 0 else "📉 Deficit"
    
    st.metric(
        label=status_label, 
        value=f"{int(abs(current_balance))} kcal",
        delta=f"{int(current_balance)} vs Burn",
        delta_color="inverse" # Red for surplus, Green for deficit
    )

# --- CONDITIONAL OURA HUD ---
# This ONLY triggers if 'oura_results' exists AND has a score above 0
oura_data = st.session_state.get("oura_results", {})

if oura_data and oura_data.get("readiness_score", 0) > 0:
    st.divider() 
    st.subheader("💍 Oura Strategic Biometrics")

    # Extracting metrics
    sleep = oura_data.get("sleep_score", 0)
    readiness = oura_data.get("readiness_score", 0)
    stress = oura_data.get("stress_level", "Stable")
    resilience = oura_data.get("resilience", "Solid")

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric("🌙 Sleep", f"{sleep}/100")
        st.progress(min(sleep / 100, 1.0) if isinstance(sleep, int) else 0.0)
        st.caption("Recovery Quality")

    with col2:
        st.metric("⚡ Readiness", f"{readiness}/100")
        st.progress(min(readiness / 100, 1.0) if isinstance(readiness, int) else 0.0)
        st.caption("Combat Ready")

    with col3:
        st.metric("📉 Stress", str(stress).capitalize())
        st.caption("System Load")

    with col4:
        st.metric("🛡️ Resilience", str(resilience).capitalize())
        st.caption("Capacity")
else:
    # If no data is found, we show absolutely nothing on the Main HUD
    pass

# ==========================================
# --- 4. THE ENGINE ---
# ==========================================

logs = load_data()
today_date = datetime.date.today()
yesterday_date = today_date - datetime.timedelta(days=1)
yesterday_str = str(yesterday_date)

if "passive_checked" not in st.session_state:
    has_yesterday_passive = not logs[(logs['date'] == yesterday_str) & (logs['type'] == 'Passive')].empty
    if not has_yesterday_passive and not logs.empty:
        save_entry("Passive", "Daily Resting Calories (BMR)", int(bmr), date_obj=yesterday_date)
        st.session_state["passive_checked"] = True
        st.rerun()
    st.session_state["passive_checked"] = True

today_logs = logs[logs['date'] == str(today_date)]
burned_exercise = abs(today_logs[today_logs['type'] == 'Exercise']['calories'].sum())
eaten = today_logs[today_logs['type'] == 'Food']['calories'].sum()
now = datetime.datetime.now()
minutes_passed = (now.hour * 60) + now.minute
passive_burn_to_now = (st.session_state["bmr"] / 1440) * minutes_passed
actual_burn_so_far = passive_burn_to_now + burned_exercise
net_mission_balance = actual_burn_so_far - eaten
total_expected_burn_today = st.session_state.get("bmr", 2000.0) + burned_exercise
remaining_budget = (total_expected_burn_today - req_deficit) - eaten

# SESSION STATE INIT
if 'df_key' not in st.session_state: st.session_state['df_key'] = 0
if 'last_workout' not in st.session_state: st.session_state['last_workout'] = None
if 'meal_desc_fill' not in st.session_state: st.session_state['meal_desc_fill'] = ""
if 'meal_cal_fill' not in st.session_state: st.session_state['meal_cal_fill'] = 0

# ==========================================
# --- 6. TABS ---
# ==========================================

tab_prt, tab_perf, tab_nut, tab_analysis = st.tabs(["🏆 PRT", "🏃 Performance", "🍽️ Nutrition", "📊 Analysis"])

# ==========================================
# --- A. PRT TAB ---
# ==========================================

with tab_prt:
    st.header("📋 Profile & Performance Assessment")
    
    # --- 1. PHYSICAL PROFILE ---
    with st.container(border=True):
        st.subheader("👤 Sailor Physical Profile")
        col1, col2 = st.columns(2)
        with col1:
            gender = st.radio("Gender", ["Male", "Female"], horizontal=True, 
                             index=0 if st.session_state.get("gender") == "Male" else 1)
            age = st.number_input("Age", value=int(st.session_state.get("age", 25)), min_value=17, max_value=65)
        
        with col2: # Line 415
            # These lines MUST be indented relative to the 'with' above
            height = st.number_input("Height (inches)", value=float(st.session_state.get("height", 70.0)), min_value=48.0, max_value=96.0, step=0.1) # Line 417
            current_weight = st.number_input("Current Weight (lbs)", value=int(st.session_state.get("current_weight", 190)), min_value=80, max_value=500)

    # --- 2. BCA GATEKEEPER ---
    maw_limit = (height * 4.2) - 105 if gender == "Male" else (height * 3.8) - 110
    over_weight = current_weight > maw_limit
    bca_fail = False

    if over_weight:
        st.warning(f"⚠️ Weight exceeds standard ({int(maw_limit)} lbs). BCA Tape required.")
        waist = st.number_input("Measured Waist (inches):", value=float(st.session_state.get("waist", 35.0)), min_value=20.0, max_value=height, step=0.1)
        whtr = waist / height
        bca_fail = whtr >= 0.55
        if bca_fail:
            st.error(f"❌ BCA FAILURE: WtHR is {whtr:.2f}")
        else:
            st.info(f"✅ BCA PASS: WtHR is {whtr:.2f}")
    else:
        st.success(f"✅ Weight within Navy standards.")

    # --- 3. PERFORMANCE SCORECARD (CLEAN MANUAL ENTRY) ---
    st.divider()
    st.subheader("🏅 PRT Scorecard")
    sc1, sc2, sc3 = st.columns(3)

    with sc1:
        st.markdown("**Upper Body**")
        pushups = st.number_input("Pushups", value=45, min_value=0, max_value=250, step=1)
        st.caption("Max reps in 2 mins")

    with sc2:
        st.markdown("**Core Stability**")
        plank_input = st.number_input("Plank (M.SS)", value=2.00, min_value=0.0, max_value=10.0, help="Example: 3.20 is 3m 20s")
        p_min = int(plank_input)
        p_sec = round((plank_input - p_min) * 100)
        plank_total_sec = (p_min * 60) + p_sec
        st.caption(f"Total: {p_min}m {p_sec}s")

    with sc3:
        st.markdown("**Cardio Event**")
        cardio_method = st.selectbox("Event Type", ["1.5 Mile Run", "Stationary Bike", "2000m Row"])
        
        # Initialize performance_val immediately
        performance_val = 0 
        c_sec = 0

        if cardio_method == "Stationary Bike":
            # Guardrails: Min 40 cals (low effort) to Max 1000 cals (Elite level 12-min burn)
            performance_val = st.number_input("Calories", value=150, min_value=40, max_value=1000, step=1)
            st.caption("Total calories burned")
            
        elif cardio_method == "2000m Row":
            # Guardrails: Min 5.00 (World Class) to Max 20.00 (Failing pace)
            cardio_input = st.number_input("Time (M.SS)", value=8.30, min_value=5.00, max_value=20.00, step=0.01)
            c_min = int(cardio_input)
            c_sec = round((cardio_input - c_min) * 100)
            performance_val = (c_min * 60) + c_sec
            st.caption(f"Total: {c_min}m {c_sec}s")
            
        else: # 1.5 Mile Run
            # Guardrails: Min 6.00 (Elite) to Max 25.00 (Walking pace)
            cardio_input = st.number_input("Time (M.SS)", value=12.30, min_value=6.00, max_value=25.00, step=0.01)
            c_min = int(cardio_input)
            c_sec = round((cardio_input - c_min) * 100)
            performance_val = (c_min * 60) + c_sec
            st.caption(f"Total: {c_min}m {c_sec}s")

    # --- 4. CALCULATION & TIER DISPLAY ---
    if p_sec < 60 and (cardio_method == "Stationary Bike" or c_sec < 60):
        s_push = calculate_score(pushups, "Pushups", age, gender)
        s_plank = calculate_score(plank_total_sec, "Plank", age, gender)
        s_cardio = calculate_score(performance_val, cardio_method, age, gender)
        
        overall_score = int((s_push + s_plank + s_cardio) / 3)
        overall_tier, overall_color = get_navy_tier(overall_score)

        st.divider()
        res_col1, res_col2, res_col3 = st.columns(3)
        res_col1.metric("Pushups", f"{s_push} pts")
        res_col2.metric("Plank", f"{s_plank} pts")
        res_col3.metric("Cardio", f"{s_cardio} pts")

        # The big visual label
        st.markdown(f"""
        <div style="background-color:rgba(255,255,255,0.05); padding:20px; border-radius:10px; border-top: 5px solid {overall_color}; text-align:center;">
            <h2 style="margin:0; color:{overall_color};">{overall_tier}</h2>
            <p style="margin:0; font-size:1.1em; opacity:0.8;">Composite Readiness Score: <strong>{overall_score}</strong></p>
        </div>
        """, unsafe_allow_html=True)

        # --- 5. LOCKDOWN VERIFICATION ---
        # Logic: If BCA fails OR any score is less than 60 (Satisfactory/Fail), Lockdown is True
        is_failing_or_sat = s_push < 60 or s_plank < 60 or s_cardio < 60
        
        if bca_fail:
            st.error("🚨 MISSION LOCKDOWN: BCA Failure.")
        elif is_failing_or_sat:
            st.warning("⚠️ MISSION LOCKDOWN: Performance below 'Good' standard. Recovery Mode Active.")

        # --- 6. SYNC BUTTON ---
        if st.button("⚖️ Update Mission Profile", type="primary", use_container_width=True):
            # Apply the lockdown state
            st.session_state["mission_lockdown"] = bca_fail or is_failing_or_sat
            
            # Update BMR for the top HUD
            w_kg, h_cm = current_weight * 0.453592, height * 2.54
            st.session_state["bmr"] = (10 * w_kg) + (6.25 * h_cm) - (5 * age) + (5 if gender == "Male" else -161)
            st.success("✅ Profile Updated. Readiness Synced.")
    else:
        st.error("🚨 Invalid Time Format: Seconds must be .00 to .59")

def run_strategic_audit():
    if not AI_READY:
        st.error(f"AI Systems Offline: {AI_ERROR}")
        return

    # Use a specific status container so the user sees progress
    with st.status("📡 Establishing Command Uplink...", expanded=True) as status:
        try:
            # 1. Gather Context
            history_data = pd.DataFrame(st.session_state.history).tail(10).to_string() if st.session_state.history else "No logs."
            goal = st.session_state.get("user_goal", "General Readiness")
            deadline = st.session_state.get("mission_deadline", datetime.date.today())
            days_rem = (deadline - datetime.date.today()).days

            prompt = f"ACT AS: Navy SEAL Fitness Specialist. GOAL: {goal}. DEADLINE: {days_rem} days. DATA: {history_data}. STRUCTURE: 1. Tactical 14-day Overwatch. 2. Strategic Phase. 3. Commander's Advice."

            # 2. Call Gemini
            status.write("🧠 Gemini 2.5 Flash analyzing trajectory...")
            plan_res = client.models.generate_content(model=MODEL_ID, contents=prompt)
            
            # 3. SAVE TO STATE
            st.session_state["master_plan"] = plan_res.text
            status.update(label="✅ Audit Complete. Trajectory Mapped.", state="complete", expanded=False)
            
        except Exception as e:
            status.update(label="❌ Audit Failed", state="error")
            st.error(f"Error: {e}")

# ==========================================
# --- B. PERFORMANCE TAB ---
# ==========================================

with tab_perf:
    st.header("⚡ Performance Command")

    # --- SECTION 1: MISSION PARAMETERS ---
    with st.container(border=True):
        st.subheader("🎯 Mission Parameters")
        col_goal, col_date = st.columns([2, 1])
        
        with col_goal:
            st.text_area("Primary Training Goal", key="user_goal", 
                         value=st.session_state.get("user_goal", "PFA Excellence"),
                         help="Detail your target scores or physical objectives.")
        
        with col_date:
            st.date_input("Mission Deadline", key="mission_deadline", 
                         value=st.session_state.get("mission_deadline", datetime.date.today() + datetime.timedelta(days=30)))

        # Logic for the AI Prompt
        today = datetime.date.today()
        deadline = st.session_state.get('mission_deadline', today)
        days_to_deadline = (deadline - today).days
        st.info(f"⏳ **{max(0, days_to_deadline)} Days** remaining until target execution.")

    # --- SECTION 2: STRATEGIC OPERATIONS (THE AI AUDIT) ---
    st.subheader("🛠️ Strategic Operations")

    # Note: Indentation is key here. This button MUST be inside 'with tab_perf'
    if st.button("🛠️ INITIALIZE STRATEGIC AUDIT", use_container_width=True, type="primary"):
        if not AI_READY:
            st.error(f"🚫 AI SYSTEMS OFFLINE: {AI_ERROR}")
        else:
            # Using a progress bar as a diagnostic tracer
            audit_progress = st.progress(10, text="Gathering mission parameters...")
            
            try:
                # 1. Gather Context
                history_data = pd.DataFrame(st.session_state.history).tail(10).to_string() if st.session_state.history else "No historical data."
                goal = st.session_state.get('user_goal', 'General Readiness')

                mission_prompt = f"""
                ACT AS: A Navy SEAL Training Specialist.
                OBJECTIVE: {goal}
                DEADLINE: {max(0, days_to_deadline)} days out.
                HISTORY: {history_data}

                INSTRUCTIONS: Provide a Day 1-14 highly detailed daily training schedule. 
                Then provide a Strategic Phase weekly theme. End with a Commander's Advice tip.
                Format the output in clean Markdown.
                """
                
                audit_progress.progress(40, text="Transmitting payload to Gemini...")
                
                # 2. Call the API
                plan_res = client.models.generate_content(model=MODEL_ID, contents=mission_prompt)
                
                audit_progress.progress(80, text="Receiving transmission...")
                
                # 3. Validate Response
                if not plan_res or not plan_res.text:
                    audit_progress.empty()
                    st.error("🚨 API responded, but the output was blank. Safety filters may have tripped.")
                else:
                    # Save and render
                    st.session_state["master_plan"] = plan_res.text
                    audit_progress.progress(100, text="✅ Trajectory Mapped & Locked to Profile.")
                    
            except Exception as e:
                audit_progress.empty()
                st.error(f"💥 CRITICAL ERROR: {str(e)}")
                st.exception(e) # Forces the full traceback to render

    # --- THE DYNAMIC DISPLAY & EXPORT HUB ---
    # This must sit entirely OUTSIDE the if st.button() block
    if "master_plan" in st.session_state:
        with st.expander("📖 VIEW ACTIVE MISSION BRIEFING", expanded=True):
            st.markdown(st.session_state["master_plan"])
            
            st.divider()
            st.caption("📦 LOGISTICS: EXPORT BRIEFING")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.download_button("📄 WORD DOC", data=create_word_doc(st.session_state["master_plan"]), 
                                   file_name="Navy_Mission_Brief.docx", use_container_width=True)
            with col2:
                st.download_button("📊 EXCEL SHEET", data=create_excel_plan(st.session_state["master_plan"]), 
                                   file_name="Tactical_Tracker.xlsx", use_container_width=True)
            with col3:
                st.download_button("📜 TEXT FILE", data=st.session_state["master_plan"], 
                                   file_name="Mission_Brief.txt", use_container_width=True)

    # --- SECTION 3: DEPLOYMENT ENVIRONMENT (GYM & GEAR) ---
    st.divider()
    with st.container(border=True):
        st.subheader("🏫 Deployment Environment")
        
        standard_gear = [
            "Bodyweight", "Standard Rack", "1.5-Mile Track", "Dumbbells", "Kettlebells", 
            "Pull-up Bar", "TRX", "Sandbags", "Rowing Machine", "Air Bike", 
            "Barbell & Plates", "Medicine Balls", "Jump Rope", "Bench", "Resistance Bands",
            "Stairmaster" # 🚨 FIX: Added so the Base Gym default doesn't crash
        ]

        if "gym_profiles" not in st.session_state:
            st.session_state.gym_profiles = {"Standard Locker": ["Bodyweight", "Standard Rack", "1.5-Mile Track"]}
        
        gym_list = list(st.session_state.gym_profiles.keys())
        active_gym = st.selectbox("Active Location", gym_list, key="active_gym_select")
        st.session_state.active_gym = active_gym

        with st.expander("🛠️ Modify Gear Locker"):
            new_name = st.text_input("New Profile Name")
            if st.button("➕ Create Profile") and new_name:
                st.session_state.gym_profiles[new_name] = ["Bodyweight"]
                st.rerun()
            
            add_item = st.selectbox("Add Standard Gear:", ["Select..."] + standard_gear)
            if st.button("📥 Add Standard") and add_item != "Select...":
                if add_item not in st.session_state.gym_profiles[active_gym]:
                    st.session_state.gym_profiles[active_gym].append(add_item)
                    st.rerun()

            st.session_state.gym_profiles[active_gym] = st.multiselect(
                "Verify Locker Contents:", 
                options=standard_gear,
                default=st.session_state.gym_profiles[active_gym]
            )

# --- SECTION 4: EVOLUTION ENGINE (DAILY WORKOUT) ---
    with st.container(border=True):
        st.subheader("🔥 Daily Evolution")
        evo_choice = st.radio("Daily Protocol:", ["Follow Strategic Plan", "Call an Audible"], horizontal=True)
        
        if st.button("🚀 Generate Daily Workout", type="primary", use_container_width=True):
            if not AI_READY:
                st.error(f"🚫 AI SYSTEMS OFFLINE: {AI_ERROR}")
            else:
                with st.spinner("Analyzing Biometrics & Generating Evolution..."):
                    try:
                        # 1. Gather the Tactical & Environmental Context
                        gear_str = ", ".join(st.session_state.gym_profiles[active_gym])
                        master_plan_context = st.session_state.get("master_plan", "General Readiness")
                        
                        # 2. Gather Real-Time Biometrics & Nutrition
                        current_eaten = int(eaten)
                        current_burn = int(active_today)
                        
                        oura_data = st.session_state.get("oura_results", {})
                        readiness = oura_data.get("readiness_score", "No Data")
                        sleep = oura_data.get("sleep_score", "No Data")
                        
                        # 3. Construct the Dynamic Prompt
                        perf_p = f"""
                        ROLE: Navy Command Fitness Leader (CFL) and Human Performance Specialist.
                        
                        LONG-TERM STRATEGY: 
                        {master_plan_context[:500]}...
                        
                        AVAILABLE GEAR: {gear_str}
                        
                        TODAY'S LIVE TELEMETRY:
                        - Calories Consumed Today: {current_eaten} kcal
                        - Active Calories Burned Today: {current_burn} kcal
                        - Oura Readiness Score: {readiness} (Out of 100)
                        - Oura Sleep Score: {sleep} (Out of 100)
                        
                        INSTRUCTIONS: 
                        Provide ONE specific workout for today that advances the LONG-TERM STRATEGY. 
                        CRITICAL: You must dynamically scale the volume and intensity based on TODAY'S LIVE TELEMETRY. 
                        - If readiness/sleep is low, or caloric deficit is extreme, pivot to active recovery, mobility, or technique work.
                        - If metrics are nominal or high, execute a standard or high-intensity evolution.
                        Briefly explain *why* the workout was scaled this way based on the data.
                        
                        Format in clean Markdown.
                        """
                        
                        # 4. Execute API Call
                        res = client.models.generate_content(model=MODEL_ID, contents=perf_p)
                        
                        # 5. Save and Render
                        if not res or not res.text:
                            st.error("🚨 API responded, but the output was blank.")
                        else:
                            st.session_state["daily_evo"] = res.text
                            st.success("✅ Evolution tailored to live biometrics.")
                            
                    except Exception as e:
                        st.error(f"💥 API Override/Failure: {str(e)}")

        # Render outside the button logic
        if "daily_evo" in st.session_state:
            st.markdown(st.session_state["daily_evo"])

    # --- SECTION 5: POST-WORKOUT ANALYSIS ---
    with st.container(border=True):
        st.subheader("📊 Performance Analysis")
        method = st.radio("Analysis:", ["AI Estimate", "Manual Tracker"], horizontal=True)
        
        if method == "AI Estimate":
            st.write("Describe your session and let the CFL estimate your caloric burn.")
            
            # 1. Add the Date Picker for historical logging
            ai_workout_date = st.date_input(
                "Workout Date", 
                value=datetime.date.today(),
                key="ai_workout_date_picker"
            )
            
            workout_desc = st.text_area("Describe the session:", placeholder="e.g., 45 mins of heavy deadlifts, plus a 2-mile run.")
            
            if st.button("🧮 Analyze & Stage", use_container_width=True):
                if not AI_READY:
                    st.error("🚫 AI SYSTEMS OFFLINE")
                elif not workout_desc:
                    st.warning("Please describe the workout first.")
                else:
                    with st.spinner("Analyzing exertion levels..."):
                        try:
                            # 2. Strict prompt formatting for easy parsing
                            prompt = f"ROLE: Navy Command Fitness Leader (CFL). Estimate the caloric burn for this workout: '{workout_desc}'. Reply ONLY in this exact format: Shortened Workout Name|TotalCalories (e.g., Heavy Leg Day & Run|650). Do not include any other text or units."
                            
                            res = client.models.generate_content(model=MODEL_ID, contents=prompt)
                            
                            if res and res.text and "|" in res.text:
                                w_name, w_cals = res.text.split("|")
                                st.session_state['ai_ex_item'] = w_name.strip()
                                st.session_state['ai_ex_cals'] = int(re.sub(r'\D', '', w_cals.strip()))
                                st.session_state['ai_ex_date'] = ai_workout_date # Save the selected date
                            else:
                                st.error("AI returned an unreadable format. Try being more specific.")
                        except Exception as e:
                            st.error(f"Analysis Failed: {str(e)}")

            # 3. AI ESTIMATE STAGING AREA
            if st.session_state.get('ai_ex_item') and st.session_state.get('ai_ex_cals'):
                st.success("Target Acquired:")
                st.info(f"**Activity:** {st.session_state['ai_ex_item']}  \n**Estimated Burn:** {st.session_state['ai_ex_cals']} kcal  \n**Date:** {st.session_state.get('ai_ex_date', datetime.date.today())}")
                
                col_c1, col_c2 = st.columns(2)
                if col_c1.button("✅ Confirm & Log", type="primary", use_container_width=True, key="ai_ex_confirm"):
                    save_entry(
                        "Exercise", 
                        st.session_state['ai_ex_item'], 
                        st.session_state['ai_ex_cals'], 
                        st.session_state['ai_ex_date']
                    )
                    # Clear staging variables
                    del st.session_state['ai_ex_item']
                    del st.session_state['ai_ex_cals']
                    del st.session_state['ai_ex_date']
                    st.toast("Workout Logged Successfully!")
                    st.rerun()
                    
                if col_c2.button("❌ Discard", use_container_width=True, key="ai_ex_discard"):
                    del st.session_state['ai_ex_item']
                    del st.session_state['ai_ex_cals']
                    if 'ai_ex_date' in st.session_state:
                        del st.session_state['ai_ex_date']
                    st.rerun()

        else:
            # Manual Tracker
            col_l1, col_l2, col_l3 = st.columns([1, 2, 1.5])
            kcal = col_l1.number_input("Kcal", min_value=0, max_value=5000, step=50)
            manual_desc = col_l2.text_input("Workout Title")
            manual_date = col_l3.date_input("Date", key="ex_manual_d")
            
            if st.button("📝 Log Work", use_container_width=True):
                if manual_desc and kcal > 0:
                    save_entry("Exercise", manual_desc, kcal, manual_date)
                    st.success(f"Logged: {manual_desc}")
                    st.rerun()

# ==========================================
# --- C. NUTRITION TAB ---
# ==========================================

with tab_nut:
    st.header("🍽️ Nutritional Command & Logistics")

    # --- 1. PERFORMANCE & GOAL SYNC ---
    # Pulling live data from other tabs for alignment
    current_goal = st.session_state.get("user_goal", "Standard Performance")
    workout_today = st.session_state.get("daily_workout", "Rest Day")
    lockdown = st.session_state.get("mission_lockdown", False)

    # --- 2. PANTRY & GYM LOCKER INVENTORY ---
    with st.expander("📦 Gym Locker & Pantry Inventory", expanded=False):
        pantry_df = load_pantry()
        
        col_p1, col_p2, col_p3 = st.columns([2, 1, 1])
        new_item = col_p1.text_input("Add Item", placeholder="e.g., Tuna, Rice", key="p_in")
        new_qty = col_p2.number_input("Qty", min_value=0, value=1, key="p_qty")
        if col_p3.button("📥 Update Stock", use_container_width=True):
            update_pantry(new_item, new_qty, action="add")
            st.rerun()

        if not pantry_df.empty:
            st.dataframe(pantry_df, use_container_width=True, hide_index=True)
            if st.button("🗑️ Purge Empty Stock"):
                pantry_df = pantry_df[pantry_df['quantity'] > 0]
                pantry_df.to_csv('pantry.csv', index=False)
                st.rerun()
        else:
            st.info("Pantry is empty. Secure supplies to optimize procurement.")

    st.divider()

    # --- 3. TACTICAL SETTINGS ---
    col_n1, col_n2 = st.columns(2)
    with col_n1:
        dining_env = st.radio("Environment:", ["⚓ Mess Decks", "🏠 Shore / Home"], horizontal=True)
        meal_freq = st.select_slider(
            "Meal Frequency:", 
            options=["1 Meal", "2 Meals", "3 Meals", "4 Meals"], 
            value="3 Meals"
        )
        duration = st.selectbox("Plan Duration:", ["Daily Brief", "14-Day Evolution"])

    with col_n2:
        zip_code = st.text_input("📍 Zip Code", placeholder="23511", max_chars=5)
        
        # 1. Pull the options dynamically from the loaded session state
        preferred_stores = st.multiselect(
            "Target Stores:",
            options=st.session_state["custom_stores"],
            default=["NEX / Commissary"] if "NEX / Commissary" in st.session_state["custom_stores"] else []
        )

        # 2. Add an expander to add new stores to the permanent roster
        with st.expander("🏪 Add Custom Store Location"):
            new_store = st.text_input("Store Name", placeholder="e.g., Whole Foods, Trader Joe's")
            if st.button("➕ Add to Roster", use_container_width=True):
                if new_store and new_store not in st.session_state["custom_stores"]:
                    st.session_state["custom_stores"].append(new_store)
                    save_stores(st.session_state["custom_stores"]) # Locks it to the JSON file
                    st.toast(f"{new_store} secured in logistics roster.")
                    st.rerun()
                elif new_store in st.session_state["custom_stores"]:
                    st.warning("Store is already in your roster.")

    # --- 4. MISSION EXECUTION (Generation & Restock) ---
    c_gen, c_restock, c_out = st.columns(3)
    
    # ACTION A: Generate the Sync'd Meal Plan
    if c_gen.button("🍴 Generate Plan", use_container_width=True, type="primary"):
        with st.spinner("Syncing with Performance Tab..."):
            stores_str = ", ".join(preferred_stores)
            nut_p = f"""
            ROLE: Navy Tactical Dietitian.
            GOAL: {current_goal}. TODAY'S WORKOUT: {workout_today}.
            DURATION: {duration}. FREQUENCY: {meal_freq}. ZIP: {zip_code}.
            ENVIRONMENT: {dining_env}.
            
            TASK: Generate a meal plan aligned with {current_goal}. 
            Adjust macros for {workout_today}. (e.g., Higher carbs for cardio, higher protein for lifting).
            """
            res = client.models.generate_content(model=MODEL_ID, contents=nut_p)
            st.session_state["active_nut_plan"] = res.text
            st.rerun()

    # ACTION B: Restock & Forecast Popover
    with c_restock.popover("🛒 Restock & Forecast", use_container_width=True):
        st.subheader("📋 Logistics Checkpoint")
        confirm_p = st.checkbox("Pantry is up-to-date")
        loadout = st.radio("Scope:", ["Match Current Plan", "Full 2-Week Strategic Loadout"])
        
        if st.button("🚀 Execute Analysis", type="primary", use_container_width=True):
            if not AI_READY:
                st.error("🚫 AI SYSTEMS OFFLINE")
            elif not confirm_p:
                st.warning("Please confirm your pantry inventory is up-to-date before forecasting.")
            else:
                with st.spinner(f"Calculating regional rates for ZIP {zip_code if zip_code else 'Unknown'}..."):
                    try:
                        inventory_str = pantry_df.to_string()
                        stores_str = ", ".join(st.session_state.get("custom_stores", preferred_stores))
                        
                        restock_p = f"""
                        ROLE: Navy Logistics Specialist. 
                        LOCATION (ZIP CODE): {zip_code if zip_code else 'General US Pricing'}
                        TARGET STORES: {stores_str}
                        SCOPE: {loadout}
                        CURRENT PANTRY INVENTORY: 
                        {inventory_str}
                        
                        TASK: 
                        1. Cross-reference the active nutrition plan with the current pantry inventory.
                        2. Provide a shopping list of missing items categorized strictly by the TARGET STORES.
                        3. Provide a hyper-local 2026 COST FORECAST for ZIP {zip_code if zip_code else 'Unknown'}. 
                           Adjust your estimated food costs to reflect exact local regional pricing for this area.
                        4. Give one actionable logistics tip for saving money at these specific stores.
                        
                        Format in clean Markdown.
                        """
                        res_r = client.models.generate_content(model=MODEL_ID, contents=restock_p)
                        st.session_state["restock_list"] = res_r.text
                        st.rerun()
                    except Exception as e:
                        st.error(f"Logistics Analysis Failed: {str(e)}")

    # ACTION C: Dining Out Recon
    with c_out.popover("🍕 Dining Out", use_container_width=True):
        st.subheader("🎯 Menu Recon")
        target_res = st.text_input("Restaurant Name", placeholder="e.g., Chipotle")
        if st.button("🔍 Get Tactical Order", use_container_width=True):
            recon_p = f"GOAL: {current_goal}. RESTAURANT: {target_res}. Give 2 orders: 1 Performance, 1 Lockdown."
            res_rec = client.models.generate_content(model=MODEL_ID, contents=recon_p)
            st.session_state["active_recon"] = res_rec.text

    # --- 5. DATA OUTPUT DISPLAY ---
    if "active_nut_plan" in st.session_state:
        st.divider()
        t1, t2 = st.tabs(["📋 Tactical Plan", "💰 Restock & Forecast"])
        with t1:
            st.markdown(st.session_state["active_nut_plan"])
        with t2:
            if "restock_list" in st.session_state:
                st.markdown(st.session_state["restock_list"])
            else:
                st.info("Run 'Restock & Forecast' to see analysis.")

    if "active_recon" in st.session_state:
        st.success("Dining Recon Complete")
        st.markdown(st.session_state["active_recon"])

    # --- 6. INTELLIGENT MEAL LOGGING ---
    st.divider()
    st.subheader("📝 Tactical Calorie Log")

    # Separate the inputs to keep the UI clean and mission-focused
    log_tab1, log_tab2, log_tab3 = st.tabs(["✍️ Manual", "🧠 AI Text Recon", "📸 AI Visual Recon"])

    # TAB 1: Classic Manual Entry
    with log_tab1:
        c1, c2, c3, c4 = st.columns([2, 1, 1.5, 1])
        q_item = c1.text_input("Item", key="ql_i", placeholder="e.g. Protein Shake")
        q_cal = c2.number_input("Kcal", value=0, step=50, key="ql_c")
        
        q_date = c3.date_input("Date", 
                           value=datetime.date.today(),
                           min_value=st.session_state.get("account_created", datetime.date(2024,1,1)),
                           max_value=datetime.date.today(),
                           key="ql_d")
                           
        if c4.button("➕ Log Manual", use_container_width=True):
            if q_item and q_cal > 0:
                save_entry("Nutrition", q_item, q_cal, q_date)
                st.toast(f"Logged {q_item} for {q_date}")
                st.rerun()
            else:
                st.error("Enter name and calories.")

    # TAB 2: Text-Based AI Estimation
    with log_tab2:
        st.write("Describe your meal in detail to estimate caloric load.")
        meal_text = st.text_area("Meal Description", placeholder="e.g., 2 eggs, 3 slices of bacon, 1 cup of black coffee")

        if st.button("🧮 Analyze Text Meal", use_container_width=True):
            if not AI_READY:
                st.error("🚫 AI SYSTEMS OFFLINE")
            elif not meal_text:
                st.warning("Enter a meal description first.")
            else:
                with st.spinner("Calculating macros..."):
                    try:
                        # Force Gemini into a strict output format for easy parsing
                        prompt = f"ROLE: Navy Dietitian. Estimate calories for this meal: '{meal_text}'. Reply ONLY in this exact format: Shortened Meal Name|TotalCalories (e.g., Eggs & Bacon Breakfast|450). Do not include any other text."
                        res = client.models.generate_content(model=MODEL_ID, contents=prompt)
                        
                        if res and res.text and "|" in res.text:
                            item, cals = res.text.split("|")
                            st.session_state['ai_log_item'] = item.strip()
                            # Use regex to strip out any accidental text characters (like 'kcal') Gemini might sneak in
                            st.session_state['ai_log_cals'] = int(re.sub(r'\D', '', cals.strip())) 
                        else:
                            st.error("AI returned an unreadable format. Try again.")
                    except Exception as e:
                        st.error(f"Analysis Failed: {str(e)}")

    # TAB 3: Visual Image Estimation
    with log_tab3:
        st.write("Provide an image of your chow tray for instant visual estimation.")
        
        # FIX: Add a 'Standby' mode so the camera doesn't auto-trigger
        img_source = st.radio(
            "Visual Input Method", 
            ["⚪ Standby", "📁 Upload", "📸 Camera"], 
            horizontal=True,
            index=0 # Forces it to start on Standby
        )
        
        meal_img = None
        if img_source == "📸 Camera":
            st.info("Ensure lighting is clear before capturing.")
            meal_img = st.camera_input("Take a picture of your food")
        elif img_source == "📁 Upload":
            meal_img = st.file_uploader("Upload meal photo", type=["jpg", "jpeg", "png"])

        if st.button("🔍 Analyze Image", use_container_width=True):
            if not AI_READY:
                st.error("🚫 AI SYSTEMS OFFLINE")
            elif img_source == "⚪ Standby":
                st.warning("Select Camera or Upload to provide an image.")
            elif not meal_img:
                st.warning("Provide an image first.")
            else:
                with st.spinner("Processing visual data..."):
                    try:
                        img = Image.open(meal_img)
                        # Ensuring the prompt uses the Navy Dietitian persona
                        prompt = "ROLE: Navy Dietitian. Estimate the total calories of the food in this image. Reply ONLY in this exact format: Shortened Meal Name|TotalCalories (e.g., Mess Deck Burger & Fries|850). Do not include any other text."
                        
                        # Gemini handles PIL images directly alongside text
                        res = client.models.generate_content(model=MODEL_ID, contents=[prompt, img])
                        
                        if res and res.text and "|" in res.text:
                            item, cals = res.text.split("|")
                            st.session_state['ai_log_item'] = item.strip()
                            st.session_state['ai_log_cals'] = int(re.sub(r'\D', '', cals.strip()))
                        else:
                            st.error("AI couldn't determine the food. Try a clearer angle.")
                    except Exception as e:
                        st.error(f"Visual Analysis Failed: {str(e)}")

    # --- AI ESTIMATE STAGING AREA ---
    # Display the staged AI entry if one exists (Shared by both AI tabs)
    if st.session_state.get('ai_log_item') and st.session_state.get('ai_log_cals'):
        st.success("Target Acquired:")
        st.info(f"**Item:** {st.session_state['ai_log_item']}  \n**Estimated Energy:** {st.session_state['ai_log_cals']} kcal")
        
        log_date = st.date_input("Assign Date to Log", value=datetime.date.today(), key="ai_log_d")
        
        col_c1, col_c2 = st.columns(2)
        if col_c1.button("✅ Confirm & Log", type="primary", use_container_width=True):
            save_entry("Nutrition", st.session_state['ai_log_item'], st.session_state['ai_log_cals'], log_date)
            # Clear staging variables
            del st.session_state['ai_log_item']
            del st.session_state['ai_log_cals']
            st.toast("Intelligence Logged Successfully!")
            st.rerun()
            
        if col_c2.button("❌ Discard", use_container_width=True):
            del st.session_state['ai_log_item']
            del st.session_state['ai_log_cals']
            st.rerun()

# ==========================================
# --- D. ANALYSIS TAB ---
# ==========================================

with tab_analysis:
    st.header("📊 Tactical Analysis & Intelligence")

    # --- SECTION 1: COMMAND SITREP ---
    # Moved from the header to the top of Analysis
    with st.container(border=True):
        st.subheader("📡 Command Sitrep")
        
        # Global metrics pulled from session state
        current_tier = st.session_state.get("current_tier", "Good Low")
        bca_pass = st.session_state.get("bca_pass", True)
        lockdown = st.session_state.get("mission_lockdown", False)
        days_rem = (st.session_state.get("mission_deadline", datetime.date.today()) - datetime.date.today()).days
        
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("Mission Proximity", f"{days_rem} Days")
        col2.metric("PRT Status", current_tier)
        col3.metric("BCA Standings", "PASS" if bca_pass else "FAIL")
        col4.metric("Readiness Level", "STRIKE" if not lockdown else "RECOVERY", delta="LOCKDOWN" if lockdown else "NOMINAL", delta_color="inverse")

    # --- SECTION 2: GENERATE INTELLIGENCE (AI AUDIT) ---
    with st.container(border=True):
        st.subheader("🕵️ Generate Intelligence")
        st.write("Synthesize all mission logs, performance trends, and nutritional data into a tactical brief.")
        
        if st.button("🛰️ Initialize Tactical Audit", use_container_width=True, type="primary"):
            if not AI_READY:
                st.error("🚫 AI SYSTEMS OFFLINE: Verify API configuration.")
            else:
                with st.spinner("Analyzing data streams..."):
                    try:
                        # Constructing the Intel Briefing
                        history_data = st.session_state.get("history", [])
                        recent_logs = history_data[-15:] if history_data else "No historical logs available."
                        
                        current_tier = st.session_state.get("current_tier", "Good Low")
                        bca_pass = st.session_state.get("bca_pass", True)
                        lockdown = st.session_state.get("mission_lockdown", False)
                        days_rem = (st.session_state.get("mission_deadline", datetime.date.today()) - datetime.date.today()).days
                        
                        intel_p = f"""
                        ROLE: Navy Command Fitness Leader (CFL) and Performance Intelligence Officer.
                        SITUATION: {days_rem} days to mission.
                        GOAL: {st.session_state.get('user_goal', 'PFA Excellence')}.
                        CURRENT STATUS: PRT {current_tier}, BCA {'Pass' if bca_pass else 'Fail'}.
                        SYSTEM STATE: {'LOCKDOWN' if lockdown else 'Optimal'}.
                        RECENT LOGS: {recent_logs}
                        
                        TASK: Generate a TACTICAL INTELLIGENCE BRIEF:
                        1. TREND ANALYSIS: Are metrics improving or stagnating?
                        2. LOGISTICAL GAP: Is current effort matching the mission goal?
                        3. IMMEDIATE ACTION: One shift in training or nutrition for the next 72 hours.
                        """
                        res = client.models.generate_content(model=MODEL_ID, contents=intel_p)
                        
                        if not res or not res.text:
                            st.error("🚨 API responded, but the output was blank.")
                        else:
                            st.session_state["tactical_intel"] = res.text
                            st.success("✅ Tactical Intelligence Brief Generated.")
                            
                    except Exception as e:
                        st.error(f"Audit Failed: {str(e)}")

        # Render output outside the button execution path
        if "tactical_intel" in st.session_state:
            st.markdown("---")
            st.markdown(st.session_state["tactical_intel"])
            if st.button("🗑️ Archive Intelligence"):
                del st.session_state["tactical_intel"]
                st.rerun()

    # --- SECTION 3: ADVANCED TELEMETRY & TRAJECTORY ---
    with st.container(border=True):
        st.subheader("📈 Logistical Trajectory")
        history_data = st.session_state.get("history", [])
        
        if history_data:
            # 1. Build DataFrame and standardize dates
            df = pd.DataFrame(history_data)
            df['Date'] = pd.to_datetime(df['Date']).dt.date
            
            # Pivot and Sum 
            daily_stats = df.groupby(['Date', 'Category'])['Value'].sum().unstack(fill_value=0)
            
            # Safety check: Ensure columns exist
            for col in ['Exercise', 'Active', 'Nutrition', 'Food']:
                if col not in daily_stats.columns:
                    daily_stats[col] = 0
                    
            # 2. Calculate the Aggregates
            daily_stats['Burned Calories'] = daily_stats['Exercise'] + daily_stats['Active']
            daily_stats['Consumed Calories'] = daily_stats['Nutrition'] + daily_stats['Food']
            
            current_bmr = st.session_state.get("bmr", 2000.0)
            daily_stats['Total_Burn'] = daily_stats['Burned Calories'] + current_bmr
            daily_stats['Net Difference'] = daily_stats['Consumed Calories'] - daily_stats['Total_Burn']
            
            # --- TACTICAL TIMELINE SELECTOR ---
            timeline = st.selectbox(
                "📅 Select Analysis Timeframe",
                ["Last 7 Days", "Last 10 Days", "Last 30 Days", "This Month", "All Time"],
                index=1
            )
            
            # Filter the dataframe based on user selection
            today = datetime.date.today()
            if timeline == "Last 7 Days":
                start_date = today - datetime.timedelta(days=7)
                filtered_df = daily_stats[daily_stats.index >= start_date]
            elif timeline == "Last 10 Days":
                start_date = today - datetime.timedelta(days=10)
                filtered_df = daily_stats[daily_stats.index >= start_date]
            elif timeline == "Last 30 Days":
                start_date = today - datetime.timedelta(days=30)
                filtered_df = daily_stats[daily_stats.index >= start_date]
            elif timeline == "This Month":
                start_date = today.replace(day=1) # First day of the current month
                filtered_df = daily_stats[daily_stats.index >= start_date]
            else:
                filtered_df = daily_stats # All Time
                
            if filtered_df.empty:
                st.warning(f"No mission telemetry found for the selected timeframe ({timeline}).")
            else:
                # --- AGGREGATE TOTALS CHART ---
                st.write(f"**📊 Aggregate Totals ({timeline})**")
                
                # Sum the filtered data
                total_burned = int(filtered_df['Burned Calories'].sum())
                total_consumed = int(filtered_df['Consumed Calories'].sum())
                total_net = int(filtered_df['Net Difference'].sum())
                
                # Create a mini dataframe just for the aggregate bar chart
                agg_df = pd.DataFrame({
                    "Category": ["🔥 Total Burned (Active)", "🍽️ Total Consumed", "⚖️ Total Net Delta"],
                    "Kcal": [total_burned, total_consumed, total_net]
                }).set_index("Category")
                
                st.bar_chart(agg_df)
                
                st.divider()

                # --- DAILY BREAKDOWN SELECTOR ---
                chart_view = st.radio(
                    f"Select Daily Breakdown ({timeline}):", 
                    ["🔥 Burned Calories", "🍽️ Consumed Calories", "⚖️ Net Difference"], 
                    horizontal=True
                )
                
                if chart_view == "🔥 Burned Calories":
                    st.write("**Daily Active Burn (Garmin + Manual Workouts)**")
                    st.line_chart(filtered_df['Burned Calories'], color="#00E5FF")
                    
                elif chart_view == "🍽️ Consumed Calories":
                    st.write("**Daily Nutritional Intake**")
                    st.line_chart(filtered_df['Consumed Calories'], color="#00C853")
                    
                else:
                    st.write("**True Net Caloric Delta (Intake vs. Total Burn + BMR)**")
                    st.caption("Values above 0 = Surplus. Values below 0 = Deficit.")
                    st.line_chart(filtered_df['Net Difference'], color="#FFAB00")
                
        else:
            # --- THE ACTIONABLE EMPTY STATE ---
            st.warning("⚠️ No mission data found.")
            
            st.markdown("""
            **To build your logistical trajectory, the system requires data input. Please execute one of the following:**
            
            * **📡 Connect a Device:** Open the **Mission Control Sidebar** on the left to uplink your Garmin, Oura, or other wearables for automatic tracking.
            * **📝 Input Workouts:** Navigate to the **Performance Tab** to manually log your physical training or use the AI to estimate exertion.
            * **🍽️ Input Meals:** Navigate to the **Nutrition Tab** to track your caloric intake via manual entry or visual AI recon.
            """)
            st.divider()

    # --- SECTION 4: DATA MAINTENANCE ---

    with st.expander("📂 Access Raw Database"):
        if history_data:
            # Display a clean, sortable dataframe
            st.dataframe(
                df.sort_values(by='Date', ascending=False), 
                use_container_width=True, 
                hide_index=True
            )
            
            st.divider()
            col_wipe1, col_wipe2 = st.columns([3, 1])
            col_wipe1.warning("Warning: This will permanently erase all local telemetry.")
            if col_wipe2.button("🔥 Purge Database", type="secondary", use_container_width=True):
                st.session_state.history = []
                if os.path.exists("fitness_log.csv"):
                    os.remove("fitness_log.csv")
                st.rerun()
        else:
            st.info("Awaiting mission data to populate trajectory charts. Start logging to build your intel baseline.")